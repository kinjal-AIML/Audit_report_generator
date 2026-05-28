import docx
import sys

def extract_text_from_paragraph(paragraph):
    return paragraph.text

def extract_text_from_table(table):
    text = []
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                text.append(paragraph.text)
    return '\n'.join(text)

def extract_text_from_docx(path):
    doc = docx.Document(path)
    full_text = []
    # paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)
    # tables
    for table in doc.tables:
        full_text.append(extract_text_from_table(table))
    # headers and footers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            full_text.append(para.text)
        for table in header.tables:
            full_text.append(extract_text_from_table(table))
        footer = section.footer
        for para in footer.paragraphs:
            full_text.append(para.text)
        for table in footer.tables:
            full_text.append(extract_text_from_table(table))
    return '\n'.join(full_text)

def find_placeholders(text, placeholders):
    results = {}
    for ph in placeholders:
        idx = text.find(ph)
        if idx != -1:
            # get surrounding context
            start = max(0, idx - 50)
            end = min(len(text), idx + len(ph) + 50)
            context = text[start:end]
            results[ph] = {
                'position': idx,
                'context': context
            }
        else:
            results[ph] = None
    return results

def main():
    template_path = r"D:\audit_report_generator\templates\UPDATED_MNSB_TEMPLATE_DYNAMIC.docx"
    try:
        text = extract_text_from_docx(template_path)
    except Exception as e:
        print(f"Error reading docx: {e}")
        sys.exit(1)
    
    placeholders = ["{{ remark }}", "{{ questions }}"]
    results = find_placeholders(text, placeholders)
    
    for ph, res in results.items():
        if res:
            print(f"Found '{ph}' at position {res['position']}")
            print(f"Context: ...{res['context']}...")
            print()
        else:
            print(f"Placeholder '{ph}' not found.")
    
    # Also optionally print all lines containing placeholders
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if "{{ remark }}" in line or "{{ questions }}" in line:
            print(f"Line {i}: {line}")

if __name__ == "__main__":
    main()