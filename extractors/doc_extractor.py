import os
from typing import List, Optional
import pandas as pd

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

# Optional Windows-only conversion from .doc -> .docx via MS Word automation
try:
    import win32com.client  # type: ignore
except Exception:  # pragma: no cover
    win32com = None


def _convert_doc_to_docx_via_word(doc_path: str) -> Optional[str]:
    if not win32com:
        return None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(doc_path))
        out_path = os.path.splitext(os.path.abspath(doc_path))[0] + ".docx"
        wdFormatXMLDocument = 12
        doc.SaveAs(out_path, FileFormat=wdFormatXMLDocument)
        doc.Close()
        word.Quit()
        return out_path
    except Exception:
        return None


essential_columns = ["Remarks", "Annexure Ref"]


def _ensure_essential_columns(df: pd.DataFrame) -> pd.DataFrame:
    for col in essential_columns:
        if col not in df.columns:
            df[col] = ""
    return df


def extract_doc_like_to_dataframe(file_path: str) -> pd.DataFrame:
    """
    Extract tables (preferred) or paragraph text from DOC/DOCX into a DataFrame.
    - For DOCX: parse with python-docx
    - For DOC: try converting to DOCX via MS Word automation (Windows), then parse
    Fallback: single-column 'Text' with paragraph lines
    """
    ext = os.path.splitext(file_path)[1].lower()

    parse_path = file_path
    if ext == ".doc":
        # Try convert to .docx
        converted = _convert_doc_to_docx_via_word(file_path)
        if converted and os.path.exists(converted):
            parse_path = converted
        else:
            # No conversion path; give up gracefully
            return pd.DataFrame(columns=["Text", "Remarks", "Annexure Ref"])  # empty placeholder

    if not Document:
        return pd.DataFrame(columns=["Text", "Remarks", "Annexure Ref"])  # python-docx not available

    try:
        doc = Document(parse_path)
        frames: List[pd.DataFrame] = []

        # Extract tables if any
        for t in doc.tables:
            rows = []
            headers = []
            # Build headers from first row if all cells have text
            if t.rows:
                first = t.rows[0].cells
                headers = [c.text.strip() or f"Col_{i}" for i, c in enumerate(first)]
            for ridx, row in enumerate(t.rows[1:] if headers else t.rows):
                cells = row.cells
                values = [c.text.strip() for c in cells]
                if not headers:
                    headers = [f"Col_{i}" for i in range(len(values))]
                rows.append(dict(zip(headers, values)))
            if rows:
                frames.append(pd.DataFrame(rows))

        # If no tables, use paragraphs as single-column text
        if not frames:
            lines = []
            for p in doc.paragraphs:
                txt = (p.text or "").strip()
                if txt:
                    lines.append({"Text": txt})
            if lines:
                frames.append(pd.DataFrame(lines))

        if not frames:
            return pd.DataFrame(columns=["Text", "Remarks", "Annexure Ref"])  # nothing extracted

        df = pd.concat(frames, ignore_index=True)
        df = df.dropna(how="all")
        return _ensure_essential_columns(df)
    except Exception:
        return pd.DataFrame(columns=["Text", "Remarks", "Annexure Ref"])  # safe fallback
