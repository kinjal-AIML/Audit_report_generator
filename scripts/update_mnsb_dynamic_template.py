import os
import shutil
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

TEMPLATES_DIR = os.path.join('templates')
TARGET = os.path.join(TEMPLATES_DIR, 'UPDATED_MNSB_TEMPLATE_DYNAMIC.docx')
BACKUP = os.path.join(TEMPLATES_DIR, 'UPDATED_MNSB_TEMPLATE_DYNAMIC.bak.docx')

def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    if bold:
        run.bold = True
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def build_template():
    doc = Document()

    # Header
    h = doc.add_paragraph('Monthly Audit Report')
    h.runs[0].bold = True
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.add_run('Branch: ').bold = True
    info.add_run('{{ branch_name }}    ')
    info.add_run('Place: ').bold = True
    info.add_run('{{ place }}    ')
    info.add_run('Period: ').bold = True
    info.add_run('{{ period_start }} to {{ period_end }}    ')
    info.add_run('Report Date: ').bold = True
    info.add_run('{{ report_date }}    ')
    info.add_run('Cash Verification: ').bold = True
    info.add_run('{{ cash_verification_date }}')

    doc.add_paragraph('')
    title = doc.add_paragraph('Questions (from Template)')
    title.runs[0].bold = True

    # Start sections loop
    doc.add_paragraph('{% for sec in sections %}')

    # Section heading
    sh = doc.add_paragraph()
    sh_run = sh.add_run('{{ sec.section_name }}')
    sh_run.bold = True

    # Table for questions
    tbl = doc.add_table(rows=1, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    set_cell_text(hdr[0], 'Sr', bold=True)
    set_cell_text(hdr[1], 'Question', bold=True)
    set_cell_text(hdr[2], 'Auditor Remark', bold=True)
    set_cell_text(hdr[3], 'Reply of Branch', bold=True)
    set_cell_text(hdr[4], 'Annexure Ref', bold=True)

    # Repeating row with Jinja tags
    row = tbl.add_row()
    c0, c1, c2, c3, c4 = row.cells
    set_cell_text(c0, '{% for q in sec.questions %}{{ q.sr_no or loop.index }}')
    set_cell_text(c1, '{{ q.question }}')
    set_cell_text(c2, '{{ answers[q.key].remark if answers.get(q.key) else "" }}')
    set_cell_text(c3, '{{ answers[q.key].branch_reply if answers.get(q.key) else "" }}')
    set_cell_text(c4, '{{ answers[q.key].annexure_reference if answers.get(q.key) else "" }}{% endfor %}')

    # End sections loop
    doc.add_paragraph('{% endfor %}')

    # Annexure summary (optional)
    doc.add_paragraph('')
    an = doc.add_paragraph('Annexure Summary:')
    an.runs[0].bold = True
    doc.add_paragraph('{{ annexure_summaries }}')

    return doc


def main():
    os.makedirs(TEMPLATES_DIR, exist_ok=True)
    if os.path.exists(TARGET):
        try:
            shutil.copyfile(TARGET, BACKUP)
        except Exception:
            pass
    doc = build_template()
    doc.save(TARGET)
    print(f"Template updated: {TARGET}")
    if os.path.exists(BACKUP):
        print(f"Backup saved: {BACKUP}")


if __name__ == '__main__':
    main()
